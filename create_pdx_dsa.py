"""
Create PDX Data Sharing Agreement - Agency Model
Based on the Experian CCDS precedent, adapted for Percayso's multi-lender agency structure
"""

from docx import Document
from docx.shared import Pt, Inches, Twips, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
import os

# Paths
source_path = r'C:\Users\DavidSant\OneDrive - Harper James Solicitors\Documents\Client Matters\Percayso\DRAFT PDX DSA HJ2.docx'
dest_path = r'C:\Users\DavidSant\OneDrive - Harper James Solicitors\Documents\Client Matters\Percayso\PDX DSA - Agency Model.docx'

# Create a new document (cleaner than trying to edit the complex original)
doc = Document()

# Set up styles
style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(10)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

# Configure Heading 1 style
h1_style = doc.styles['Heading 1']
h1_style.font.name = 'Arial'
h1_style.font.size = Pt(14)
h1_style.font.bold = True
h1_style.font.color.rgb = None  # Black
h1_style.paragraph_format.space_before = Pt(18)
h1_style.paragraph_format.space_after = Pt(6)

# Configure Heading 2 style
h2_style = doc.styles['Heading 2']
h2_style.font.name = 'Arial'
h2_style.font.size = Pt(12)
h2_style.font.bold = True
h2_style.font.color.rgb = None
h2_style.paragraph_format.space_before = Pt(12)
h2_style.paragraph_format.space_after = Pt(6)

# ============================================================================
# MULTILEVEL LIST SETUP
# ============================================================================

def create_multilevel_numbering(doc):
    """Create a multilevel numbering definition for legal documents"""
    # Get or create numbering part
    numbering_part = doc.part.numbering_part
    
    # Get the numbering element
    numbering = numbering_part.numbering_definitions._numbering
    
    # Create abstract numbering
    abstractNumId = 1
    
    # Create abstractNum element
    abstractNum = OxmlElement('w:abstractNum')
    abstractNum.set(qn('w:abstractNumId'), str(abstractNumId))
    
    # Multi-level type
    multiLevelType = OxmlElement('w:multiLevelType')
    multiLevelType.set(qn('w:val'), 'multilevel')
    abstractNum.append(multiLevelType)
    
    # Level 0: 1. 2. 3. (Clause numbers)
    lvl0 = OxmlElement('w:lvl')
    lvl0.set(qn('w:ilvl'), '0')
    
    start0 = OxmlElement('w:start')
    start0.set(qn('w:val'), '1')
    lvl0.append(start0)
    
    numFmt0 = OxmlElement('w:numFmt')
    numFmt0.set(qn('w:val'), 'decimal')
    lvl0.append(numFmt0)
    
    lvlText0 = OxmlElement('w:lvlText')
    lvlText0.set(qn('w:val'), '%1.')
    lvl0.append(lvlText0)
    
    lvlJc0 = OxmlElement('w:lvlJc')
    lvlJc0.set(qn('w:val'), 'left')
    lvl0.append(lvlJc0)
    
    pPr0 = OxmlElement('w:pPr')
    ind0 = OxmlElement('w:ind')
    ind0.set(qn('w:left'), '720')
    ind0.set(qn('w:hanging'), '720')
    pPr0.append(ind0)
    lvl0.append(pPr0)
    
    rPr0 = OxmlElement('w:rPr')
    b0 = OxmlElement('w:b')
    rPr0.append(b0)
    lvl0.append(rPr0)
    
    abstractNum.append(lvl0)
    
    # Level 1: 1.1, 1.2, 1.3 (Sub-clauses)
    lvl1 = OxmlElement('w:lvl')
    lvl1.set(qn('w:ilvl'), '1')
    
    start1 = OxmlElement('w:start')
    start1.set(qn('w:val'), '1')
    lvl1.append(start1)
    
    numFmt1 = OxmlElement('w:numFmt')
    numFmt1.set(qn('w:val'), 'decimal')
    lvl1.append(numFmt1)
    
    lvlText1 = OxmlElement('w:lvlText')
    lvlText1.set(qn('w:val'), '%1.%2')
    lvl1.append(lvlText1)
    
    lvlJc1 = OxmlElement('w:lvlJc')
    lvlJc1.set(qn('w:val'), 'left')
    lvl1.append(lvlJc1)
    
    pPr1 = OxmlElement('w:pPr')
    ind1 = OxmlElement('w:ind')
    ind1.set(qn('w:left'), '720')
    ind1.set(qn('w:hanging'), '720')
    pPr1.append(ind1)
    lvl1.append(pPr1)
    
    abstractNum.append(lvl1)
    
    # Level 2: (a), (b), (c) (List items)
    lvl2 = OxmlElement('w:lvl')
    lvl2.set(qn('w:ilvl'), '2')
    
    start2 = OxmlElement('w:start')
    start2.set(qn('w:val'), '1')
    lvl2.append(start2)
    
    numFmt2 = OxmlElement('w:numFmt')
    numFmt2.set(qn('w:val'), 'lowerLetter')
    lvl2.append(numFmt2)
    
    lvlText2 = OxmlElement('w:lvlText')
    lvlText2.set(qn('w:val'), '(%3)')
    lvl2.append(lvlText2)
    
    lvlJc2 = OxmlElement('w:lvlJc')
    lvlJc2.set(qn('w:val'), 'left')
    lvl2.append(lvlJc2)
    
    pPr2 = OxmlElement('w:pPr')
    ind2 = OxmlElement('w:ind')
    ind2.set(qn('w:left'), '1440')
    ind2.set(qn('w:hanging'), '720')
    pPr2.append(ind2)
    lvl2.append(pPr2)
    
    abstractNum.append(lvl2)
    
    # Level 3: (i), (ii), (iii) (Sub-list items)
    lvl3 = OxmlElement('w:lvl')
    lvl3.set(qn('w:ilvl'), '3')
    
    start3 = OxmlElement('w:start')
    start3.set(qn('w:val'), '1')
    lvl3.append(start3)
    
    numFmt3 = OxmlElement('w:numFmt')
    numFmt3.set(qn('w:val'), 'lowerRoman')
    lvl3.append(numFmt3)
    
    lvlText3 = OxmlElement('w:lvlText')
    lvlText3.set(qn('w:val'), '(%4)')
    lvl3.append(lvlText3)
    
    lvlJc3 = OxmlElement('w:lvlJc')
    lvlJc3.set(qn('w:val'), 'left')
    lvl3.append(lvlJc3)
    
    pPr3 = OxmlElement('w:pPr')
    ind3 = OxmlElement('w:ind')
    ind3.set(qn('w:left'), '2160')
    ind3.set(qn('w:hanging'), '720')
    pPr3.append(ind3)
    lvl3.append(pPr3)
    
    abstractNum.append(lvl3)
    
    # Insert abstractNum at the beginning
    numbering.insert(0, abstractNum)
    
    # Create num element that references the abstractNum
    num = OxmlElement('w:num')
    num.set(qn('w:numId'), '1')
    abstractNumId_ref = OxmlElement('w:abstractNumId')
    abstractNumId_ref.set(qn('w:val'), str(abstractNumId))
    num.append(abstractNumId_ref)
    numbering.append(num)
    
    return 1  # Return numId

def apply_numbering(paragraph, numId, level):
    """Apply numbering to a paragraph"""
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    
    numPr = OxmlElement('w:numPr')
    ilvl = OxmlElement('w:ilvl')
    ilvl.set(qn('w:val'), str(level))
    numPr.append(ilvl)
    
    numId_elem = OxmlElement('w:numId')
    numId_elem.set(qn('w:val'), str(numId))
    numPr.append(numId_elem)
    
    pPr.insert(0, numPr)

# Create the numbering definition
numId = create_multilevel_numbering(doc)

# Helper function to add a clause heading (Level 0)
def add_clause_heading(doc, title):
    p = doc.add_paragraph()
    p.style = 'Normal'
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(12)
    apply_numbering(p, numId, 0)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    return p

# Helper function to add a sub-clause (Level 1)
def add_sub_clause(doc, text):
    p = doc.add_paragraph()
    p.style = 'Normal'
    run = p.add_run(text)
    apply_numbering(p, numId, 1)
    return p

# Helper function to add a list item (a), (b), etc. (Level 2)
def add_list_item(doc, text):
    p = doc.add_paragraph()
    p.style = 'Normal'
    run = p.add_run(text)
    apply_numbering(p, numId, 2)
    return p

# Helper function for bullet points
def add_bullet(doc, text, indent_level=1):
    p = doc.add_paragraph()
    p.style = 'List Bullet'
    run = p.add_run(text)
    return p

# ============================================================================
# COVER LETTER
# ============================================================================

title = doc.add_paragraph()
title.style = 'Title'
title_run = title.add_run("DATA SHARING AGREEMENT")
title_run.bold = True
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run("BETWEEN:").bold = True
doc.add_paragraph()

doc.add_paragraph("(1)\tPERCAYSO LIMITED (Company No. [NUMBER]) whose registered office is at Hine House, 25 Regent Street, Nottingham, NG1 5BS (\"Percayso\"), acting as agent for and on behalf of the Participating Lenders listed in Schedule 1; and")
doc.add_paragraph()
doc.add_paragraph("(2)\t[CRA NAME] (Company No. [NUMBER]) whose registered office is at [ADDRESS] (the \"CRA\"), being a designated credit reference agency under the Regulations.")
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run("BACKGROUND").bold = True
doc.add_paragraph()

background_text = [
    "(A)\tPercayso operates a data sharing platform known as \"PDX\" that facilitates the sharing of credit information between Finance Providers, Credit Providers and designated credit reference agencies pursuant to the Regulations.",
    "(B)\tEach Participating Lender has entered into a Lender Agreement with Percayso, under which (among other things) the Participating Lender has appointed Percayso as its agent to enter into data sharing arrangements with designated CRAs on its behalf.",
    "(C)\tThe CRA wishes to receive Credit Information from the Participating Lenders via the PDX Platform in accordance with the Regulations.",
    "(D)\tPercayso enters into this Agreement as agent for and on behalf of each Participating Lender, and each Participating Lender shall be bound by the terms of this Agreement as if it were a party to it."
]

for para in background_text:
    doc.add_paragraph(para)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("AGREED TERMS").bold = True
doc.add_paragraph()

# ============================================================================
# CLAUSE 1: DEFINITIONS
# ============================================================================

add_clause_heading(doc, "DEFINITIONS AND INTERPRETATION")

add_sub_clause(doc, "In this Agreement the following words and expressions shall have the following meanings:")

definitions = [
    "\"Agreement\" means this data sharing agreement, including all Schedules;",
    "\"Applicable Laws\" means all legislation, regulations, and other rules having equivalent force which are applicable to that party, including the Regulations, the Data Protection Act 2018, and the UK GDPR;",
    "\"Confidential Information\" means any and all information relating to the trade secrets, operations, processes, plans, product information, know-how, customer lists, transactions, affairs and/or business of the parties and/or their customers, suppliers, clients or group companies;",
    "\"CRA\" means the credit reference agency identified in this Agreement, being a designated credit reference agency under the Regulations;",
    "\"Credit Information\" means the credit information which Participating Lenders make available to the CRA pursuant to this Agreement (as set out in Schedule 2);",
    "\"Credit Provider\" means an organisation that offers credit as a result of a service it provides to a business;",
    "\"Effective Date\" means the date of this Agreement;",
    "\"Finance Provider\" shall be interpreted in accordance with the Regulations and includes a bank designated under the Regulations;",
    "\"Group Company\" means any company which is a subsidiary, holding company or subsidiary of a holding company (as defined by section 1159 of the Companies Act 2006);",
    "\"Intellectual Property Rights\" means copyright, database right, patents, registered and unregistered design rights, registered and unregistered trademarks, and all other industrial, commercial or intellectual property rights existing in any jurisdiction;",
    "\"Lender Agreement\" means the agreement between Percayso and a Participating Lender under which Percayso provides PDX Platform services and is appointed as agent to enter into data sharing arrangements on behalf of that Participating Lender;",
    "\"Participating Lender\" means a Finance Provider or Credit Provider that has entered into a Lender Agreement with Percayso, is listed in Schedule 1, and whose participation has not been terminated;",
    "\"PDX Platform\" means Percayso's data validation, transformation and submission platform;",
    "\"Percayso\" means Percayso Limited, acting as agent for the Participating Lenders;",
    "\"Principles of Reciprocity\" means the Principles of Reciprocity as published by the Steering Committee on Reciprocity from time to time;",
    "\"Regulations\" means the Small and Medium Sized Business (Credit Information) Regulations 2015 as amended or replaced from time to time;",
    "\"UK GDPR\" means the General Data Protection Regulation (EU) 2016/679 as it forms part of domestic law in the United Kingdom by virtue of section 3 of the European Union (Withdrawal) Act 2018."
]

for defn in definitions:
    add_list_item(doc, defn)

add_sub_clause(doc, "In this Agreement:")

interp_rules = [
    "any reference to a statutory provision includes a reference to any modification or re-enactment of it from time to time;",
    "references to clauses and schedules are to the clauses of and the schedules to this Agreement;",
    "the singular includes the plural and vice versa;",
    "headings are for ease of reference only and shall not affect the construction or interpretation of this Agreement;",
    "where any matter is to be agreed, such agreement must be recorded in writing to take effect;",
    "wherever the words \"including\", \"include\", \"includes\" or \"included\" are used they shall be deemed to be followed by the words \"without limitation\"."
]
for rule in interp_rules:
    add_list_item(doc, rule)

# ============================================================================
# CLAUSE 2: AGENCY
# ============================================================================

add_clause_heading(doc, "AGENCY")

add_sub_clause(doc, "Percayso enters into this Agreement as agent for and on behalf of each Participating Lender. Each Participating Lender shall be bound by the terms of this Agreement as if it were a party to it.")
add_sub_clause(doc, "Percayso warrants that it has been duly authorised by each Participating Lender to enter into this Agreement on its behalf, and to bind the Participating Lender to the obligations set out herein.")
add_sub_clause(doc, "Schedule 1 sets out the current list of Participating Lenders as at the Effective Date. Percayso may add or remove Participating Lenders by giving written notice to the CRA in accordance with Clause 2.4.")
add_sub_clause(doc, "To add a Participating Lender, Percayso shall give the CRA not less than 10 Business Days' written notice, specifying the name, registered number and registered address of the new Participating Lender, and the date from which it shall become bound by this Agreement. The new Participating Lender shall become bound by this Agreement on the date specified in such notice.")
add_sub_clause(doc, "To remove a Participating Lender, Percayso shall give the CRA not less than 30 days' written notice. Removal shall take effect on the date specified in such notice, subject to any outstanding obligations of the Participating Lender under this Agreement.")
add_sub_clause(doc, "The CRA may request that Percayso remove a Participating Lender if the CRA reasonably believes that the Participating Lender is in material breach of this Agreement. Percayso shall respond to such request within 10 Business Days and, if the CRA's concerns are well-founded, shall take reasonable steps to procure that the breach is remedied or, failing that, shall remove the Participating Lender.")
add_sub_clause(doc, "For the avoidance of doubt, Percayso is not itself a data supplier, Finance Provider or Credit Provider under the Regulations. The data sharing relationship is between each Participating Lender and the CRA, with Percayso acting as agent and service provider.")

# ============================================================================
# CLAUSE 3: SUPPLY OF CREDIT INFORMATION
# ============================================================================

add_clause_heading(doc, "SUPPLY OF CREDIT INFORMATION")

add_sub_clause(doc, "Subject to paragraph 6(2) of the Regulations, on the Effective Date (or, in the case of a Participating Lender added after the Effective Date, on the date specified in the notice of addition), each Participating Lender shall make the Credit Information available to the CRA via the PDX Platform for the purposes set out in Clause 4, and shall provide updates to the Credit Information on a monthly basis thereafter.")
add_sub_clause(doc, "The Credit Information shall be provided in a format agreed between Percayso and the CRA from time to time. Percayso shall validate and transform the Credit Information received from Participating Lenders to comply with the CRA's format requirements before submission to the CRA.")
add_sub_clause(doc, "Where testing of the Credit Information is required prior to first submission by a Participating Lender, Percayso shall coordinate such testing with the CRA. The CRA agrees that it shall only use test data solely for the purpose of the agreed tests and will not include test data within its live credit bureau.")
add_sub_clause(doc, "Each Participating Lender shall use reasonable endeavours to ensure that the Credit Information it provides is true and accurate. The CRA acknowledges that Percayso performs validation services but does not warrant the accuracy of Credit Information provided by Participating Lenders.")
add_sub_clause(doc, "If a Participating Lender becomes aware that it may have supplied inaccurate or misleading Credit Information, it shall (via Percayso) notify the CRA as soon as reasonably practicable and take all reasonable steps to correct the Credit Information.")

# ============================================================================
# CLAUSE 4: LICENCE
# ============================================================================

add_clause_heading(doc, "LICENCE")

add_sub_clause(doc, "Each Participating Lender grants to the CRA a non-exclusive, non-transferable, royalty-free licence to use its Credit Information:")

licence_purposes = [
    "to develop and sell products and services to assist Finance Providers and Credit Providers in assessing whether to offer a business finance, lending or credit, and the ongoing management of such finance, lending or credit, including for identity verification, fraud prevention, assessing creditworthiness, debt collection, and regulatory compliance;",
    "to develop and sell scores to Trade Credit Providers for similar purposes, provided the CRA requires the Trade Credit Provider to enter into an appropriate agreement and the business has been notified;",
    "for the purposes necessary for the CRA to comply with its obligations under the Regulations;",
    "to provide a copy of the Credit Information relating to any business to that business on request;",
    "for such other purposes as are consistent with the Principles of Reciprocity or as agreed between the parties from time to time."
]
for purpose in licence_purposes:
    add_list_item(doc, purpose)

add_sub_clause(doc, "The licence to use the Credit Information shall expire on termination of a Participating Lender's participation in this Agreement, save that the CRA may retain such Credit Information to the extent required to comply with the Regulations or any other Applicable Law.")

# ============================================================================
# CLAUSE 5: FLOW-DOWN WARRANTIES
# ============================================================================

add_clause_heading(doc, "FLOW-DOWN WARRANTIES")

add_sub_clause(doc, "Percayso warrants that, under each Lender Agreement, the Participating Lender has warranted and agreed that:")

warranties = [
    "it is a Finance Provider or Credit Provider for the purposes of the Regulations;",
    "each customer in respect of whom it provides Credit Information has been notified of, and has agreed to, the Credit Information being provided to designated CRAs in accordance with the Regulations;",
    "it will use reasonable endeavours to ensure that the Credit Information it provides is true and accurate;",
    "it will notify Percayso promptly if it becomes aware that it may have supplied inaccurate or misleading Credit Information, and will take reasonable steps to correct such information;",
    "it will comply with all Applicable Laws, including Data Protection Laws, in connection with its provision of Credit Information;",
    "it has appointed Percayso as its agent to enter into data sharing arrangements with designated CRAs on its behalf, and agrees to be bound by the terms of any such arrangements."
]
for warranty in warranties:
    add_list_item(doc, warranty)

add_sub_clause(doc, "Percayso shall ensure that each Lender Agreement contains warranties and obligations from the Participating Lender that are no less onerous than those set out in Clause 5.1.")
add_sub_clause(doc, "On reasonable request, Percayso shall provide the CRA with a copy of the template Lender Agreement (or relevant extracts) to confirm compliance with Clause 5.2.")
add_sub_clause(doc, "At the CRA's request, Percayso shall use reasonable endeavours to procure that a Participating Lender provides a copy of the notification provided to its customers from time to time regarding the sharing of Credit Information.")

# ============================================================================
# CLAUSE 6: COMPLIANCE AND DATA PROTECTION
# ============================================================================

add_clause_heading(doc, "COMPLIANCE AND DATA PROTECTION")

add_sub_clause(doc, "Each party undertakes to the other that, in connection with this Agreement, it will at all times comply with all relevant Applicable Laws.")
add_sub_clause(doc, "To the extent any Credit Information is personal data, each Participating Lender is the data controller of its Credit Information. The CRA will become a data controller of such personal data when it is received from the Participating Lender via the PDX Platform.")
add_sub_clause(doc, "Percayso acts as a data processor on behalf of each Participating Lender in respect of the validation, transformation and transmission of Credit Information via the PDX Platform. Percayso's processing activities are governed by the Lender Agreement with each Participating Lender.")
add_sub_clause(doc, "The CRA warrants that it shall take appropriate technical and organisational measures against unauthorised or unlawful processing of the Credit Information and against accidental loss or destruction of, or damage to, the Credit Information.")
add_sub_clause(doc, "The CRA shall permit Percayso (on behalf of the Participating Lenders), on reasonable notice, during normal working hours and up to a maximum of once per calendar year, to audit its compliance with its obligations under this Agreement.")
add_sub_clause(doc, "Each party agrees to promptly provide any information and assistance reasonably requested by the other party in order to respond to any query or complaint relating to the Credit Information, including to comply with data subject access requests or requirements under the Regulations.")

# ============================================================================
# CLAUSE 7: LIABILITY
# ============================================================================

add_clause_heading(doc, "LIABILITY")

add_sub_clause(doc, "Neither party shall be liable for indirect or consequential loss suffered or incurred by the other party.")
add_sub_clause(doc, "Nothing in this Agreement shall operate to exclude or limit either party's liability for death or personal injury caused by negligence, for fraud or fraudulent misrepresentation, or for any other matter which it would be illegal for that party to exclude or limit its liability.")
add_sub_clause(doc, "Percayso shall not be liable to the CRA for any breach of this Agreement by a Participating Lender, save to the extent that such breach arises from Percayso's own negligence or wilful default.")
add_sub_clause(doc, "Each Participating Lender shall be liable for its own breaches of this Agreement.")

# ============================================================================
# CLAUSE 8: TERM AND TERMINATION
# ============================================================================

add_clause_heading(doc, "TERM AND TERMINATION")

add_sub_clause(doc, "This Agreement shall begin on the Effective Date and, subject to the termination rights set out below, shall remain in force until terminated in accordance with this Clause 8.")
add_sub_clause(doc, "Either party may terminate this Agreement at any time on giving the other not less than 90 days' written notice.")
add_sub_clause(doc, "Either party shall be entitled to terminate this Agreement immediately by serving written notice on the other party if:")

add_list_item(doc, "the CRA's designation as a credit reference agency under the Regulations is revoked; or")
add_list_item(doc, "Percayso ceases to operate the PDX Platform.")

add_sub_clause(doc, "Either party shall be entitled to terminate this Agreement immediately by serving written notice on the other party if:")

add_list_item(doc, "the other party commits a material breach of any of its obligations under this Agreement which is not capable of remedy;")
add_list_item(doc, "the other party commits a material breach of any of its obligations under this Agreement which is not remedied within 28 days after receipt of a notice specifying the breach and requiring its remedy;")
add_list_item(doc, "the other party has passed a resolution for its winding up (save for voluntary reconstruction or amalgamation), is subject to administration or receivership, is unable to pay its debts within the meaning of section 123 Insolvency Act 1986, or ceases to trade.")

add_sub_clause(doc, "Termination of a Participating Lender's participation (whether by removal under Clause 2.5 or 2.6, or because the Participating Lender ceases to be a Finance Provider or Credit Provider) shall not affect this Agreement as between Percayso (on behalf of the remaining Participating Lenders) and the CRA.")
add_sub_clause(doc, "Upon termination of this Agreement, the CRA may retain Credit Information to the extent required to comply with the Regulations or any other Applicable Law, after which time the CRA shall delete such Credit Information from its databases.")
add_sub_clause(doc, "Termination of this Agreement shall not affect any rights, obligations or liabilities of either party which have accrued before termination or which are intended to continue beyond termination.")

# ============================================================================
# CLAUSE 9: INTELLECTUAL PROPERTY
# ============================================================================

add_clause_heading(doc, "INTELLECTUAL PROPERTY RIGHTS")

add_sub_clause(doc, "Any Intellectual Property Rights in the Credit Information shall remain the property of the relevant Participating Lender.")
add_sub_clause(doc, "All Intellectual Property Rights in the databases containing Credit Information (but excluding the Credit Information itself) and any products or services provided by the CRA using the Credit Information shall remain the property of the CRA.")
add_sub_clause(doc, "All Intellectual Property Rights in the PDX Platform shall remain the property of Percayso.")

# ============================================================================
# CLAUSE 10: CONFIDENTIALITY
# ============================================================================

add_clause_heading(doc, "CONFIDENTIALITY")

add_sub_clause(doc, "Each party shall keep the Confidential Information of the other party strictly confidential and not disclose it to any person except as permitted by this Agreement.")
add_sub_clause(doc, "Each party may disclose Confidential Information to employees, officers, auditors and professional advisors who require it for the performance of this Agreement, provided that such persons are bound by equivalent confidentiality obligations.")
add_sub_clause(doc, "The restrictions in Clause 10.1 do not apply to information that:")

conf_exceptions = [
    "is or comes within the public domain other than through breach of this Agreement;",
    "was in the recipient's possession before receiving it from the other party;",
    "is lawfully received from a third party;",
    "is independently developed without use of Confidential Information; or",
    "is required to be disclosed by law or by a court of competent jurisdiction."
]
for exception in conf_exceptions:
    add_list_item(doc, exception)

add_sub_clause(doc, "The CRA shall not identify any Participating Lender as the source of any Credit Information, except as required by Applicable Law or to respond to a request by a business for its own Credit Information.")
add_sub_clause(doc, "This Clause 10 shall survive termination of this Agreement.")

# ============================================================================
# CLAUSE 11: GENERAL
# ============================================================================

add_clause_heading(doc, "GENERAL")

add_sub_clause(doc, "Any notices under this Agreement shall be in writing and shall be delivered personally or sent by recorded delivery to the registered office of each party, or as otherwise notified. Notices shall be deemed given on delivery (if delivered personally) or two clear days after posting (if sent by post).")
add_sub_clause(doc, "Neither party may assign, transfer, charge or deal in any other manner with this Agreement or any of its rights under it without the prior written consent of the other party (such consent not to be unreasonably withheld or delayed).")
add_sub_clause(doc, "Except as expressly provided, nothing in this Agreement shall create rights enforceable by any person who is not a party to this Agreement.")
add_sub_clause(doc, "This Agreement sets out all the terms agreed between the parties relating to its subject matter and supersedes any previous agreement between the parties relating to the same subject matter.")
add_sub_clause(doc, "Neither party will be liable for any delay or failure in performance due to circumstances beyond its reasonable control.")
add_sub_clause(doc, "This Agreement shall be governed by the laws of England and Wales. The English courts shall have exclusive jurisdiction over any claim arising out of or in connection with this Agreement.")

# ============================================================================
# SIGNATURE BLOCK
# ============================================================================

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph("SIGNED by the parties on the date first written above.")
doc.add_paragraph()
doc.add_paragraph()

sig_block = """
Signed for and on behalf of                    Signed for and on behalf of
PERCAYSO LIMITED                               [CRA NAME]
(as agent for the Participating Lenders)       


_______________________________                _______________________________
Signature                                      Signature


_______________________________                _______________________________
Name                                           Name


_______________________________                _______________________________
Date                                           Date
"""
doc.add_paragraph(sig_block)

# ============================================================================
# SCHEDULE 1: PARTICIPATING LENDERS
# ============================================================================

doc.add_page_break()

sched1_title = doc.add_paragraph()
sched1_title.style = 'Heading 1'
sched1_title.add_run("SCHEDULE 1")

sched1_subtitle = doc.add_paragraph()
sched1_subtitle.style = 'Heading 2'
sched1_subtitle.add_run("PARTICIPATING LENDERS")

doc.add_paragraph()
doc.add_paragraph("The following Finance Providers and Credit Providers are Participating Lenders as at the Effective Date:")
doc.add_paragraph()

# Create a proper table
table = doc.add_table(rows=4, cols=3)
table.style = 'Table Grid'

# Header row
header_cells = table.rows[0].cells
header_cells[0].text = "Name"
header_cells[1].text = "Company Number"
header_cells[2].text = "Effective Date"

# Make header bold
for cell in header_cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True

# Sample rows
table.rows[1].cells[0].text = "[Lender 1 Name]"
table.rows[1].cells[1].text = "[Number]"
table.rows[1].cells[2].text = "[Date]"

table.rows[2].cells[0].text = "[Lender 2 Name]"
table.rows[2].cells[1].text = "[Number]"
table.rows[2].cells[2].text = "[Date]"

table.rows[3].cells[0].text = "[Lender 3 Name]"
table.rows[3].cells[1].text = "[Number]"
table.rows[3].cells[2].text = "[Date]"

doc.add_paragraph()
doc.add_paragraph("[Additional lenders to be added by notice in accordance with Clause 2.4]")

# ============================================================================
# SCHEDULE 2: CREDIT INFORMATION
# ============================================================================

doc.add_page_break()

sched2_title = doc.add_paragraph()
sched2_title.style = 'Heading 1'
sched2_title.add_run("SCHEDULE 2")

sched2_subtitle = doc.add_paragraph()
sched2_subtitle.style = 'Heading 2'
sched2_subtitle.add_run("CREDIT INFORMATION")

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run("Information relating to a loan made in sterling to the business:").bold = True

loan_info = [
    "start date of loan agreement;",
    "the date the loan is due to be fully repaid, has been fully repaid or enters default;",
    "amount of loan outstanding;",
    "repayment period;",
    "repayment frequency;",
    "repayment amount;",
    "number of missed payments;",
    "details of any defaults and associated satisfactions."
]
for item in loan_info:
    add_bullet(doc, item)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("Information relating to a credit card account denominated in sterling and held in the name of the business:").bold = True

card_info = [
    "start date of the facility;",
    "the date the facility closed (if applicable);",
    "outstanding balance;",
    "agreed credit limit;",
    "number of missed payments;",
    "number of cash advances;",
    "value of cash advances;",
    "details of any defaults and associated satisfactions."
]
for item in card_info:
    add_bullet(doc, item)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("Information relating to a current account denominated in sterling and held in the name of the business:").bold = True

account_info = [
    "start date of the facility;",
    "the date the facility closed (if applicable);",
    "current balance;",
    "minimum balance;",
    "maximum balance;",
    "average balance;",
    "overdraft limit;",
    "total value of all payments into the account;",
    "total value of debits withdrawn from the account;",
    "number of days in month where the customer has exceeded its approved limit;",
    "number of cheques or direct debits that have not been paid due to insufficient funds."
]
for item in account_info:
    add_bullet(doc, item)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("Where any of the information described above is provided, the following shall also be provided:").bold = True

business_info = [
    "business type indicator (e.g. limited liability company or non-limited business);",
    "business name and address;",
    "company registration number (if applicable);",
    "telephone number;",
    "VAT number (if applicable)."
]
for item in business_info:
    add_bullet(doc, item)

# ============================================================================
# SAVE THE DOCUMENT
# ============================================================================

doc.save(dest_path)
print(f"Document saved to: {dest_path}")
print("Done!")
